#!/usr/bin/env python3
"""Generate docx稿子 from cs-core and en-interview pack content. Run: pip install python-docx && python build_docx.py"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_question_block(doc, q, pack_name):
    doc.add_paragraph(q["title"], style="Heading 2")
    for i, v in enumerate(q["variants"], 1):
        text = v if isinstance(v, str) else v.get("text", "")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        for line in text.strip().split("\n"):
            p.add_run(line + "\n")
    doc.add_paragraph()  # gap after question

def main():
    # --- cs-core ---
    cs_core = [
        {
            "id": 1,
            "title": "1. Spiral system design",
            "variants": [
                "10s\n\nI built a system called Spiral that explores how event histories can be represented structurally instead of relying on mutable state.\n\nThe main challenge was designing an execution model that remained simple while still being practical to use.",
                "20s\n\nOne project I've been working on is called Spiral.\n\nIt explores how software systems can represent event histories using an append-only structure instead of traditional mutable state.\n\nThe main challenge was designing an execution model that kept the system understandable while still being practical to implement.",
                "40s\n\nOne project I've spent a lot of time on is called Spiral.\n\nThe idea started from a question about how software systems represent history and state.\n\nInstead of mutating state directly, Spiral records events in an append-only structure and derives the current state from their relationships.\n\nEarly on the architecture became too complex, so I simplified the model and focused on the core execution rules.\n\nLater I built a prototype platform so people could interact with the system.",
                "90s\n\nOne project I've spent a lot of time on is called Spiral.\n\nThe project started from a conceptual question: how can software systems represent event histories in a way that's easier to reason about?\n\nTraditional systems rely heavily on mutable state, which can make it difficult to understand how the system reached a particular condition.\n\nIn Spiral I experimented with an append-only event structure, where every change becomes a new event and the system derives the current state from relationships between those events.\n\nEarly in development the architecture became overly complex because I tried to support too many features.\n\nEventually I simplified the design and focused on the core execution model.\n\nAfter that I built a prototype platform where users could interact with the system through a structured editor.\n\nWhat I enjoyed most was connecting system architecture with a usable interface.",
            ],
        },
        {
            "id": 2,
            "title": "2. Robotics teamwork story",
            "variants": [
                "10s\n\nDuring a robotics competition a wheel came off during a match.\n\nOur team quickly divided tasks and replaced it so we could continue.",
                "20s\n\nDuring a robotics competition I worked on a small team responsible for maintaining the robot.\n\nIn one match a wheel came off, so we had to react quickly and divide tasks to repair it.",
                "40s\n\nDuring a robotics competition I worked on a small team responsible for maintaining the robot.\n\nIn one match a wheel came off during operation.\n\nWe reacted quickly and divided tasks immediately.\n\nOne teammate stabilized the robot while another teammate and I replaced the wheel assembly.\n\nIt showed me how important communication and coordination are in a team.",
                "90s\n\nDuring a robotics competition I worked as part of a small team responsible for building and maintaining the robot.\n\nDuring one of the matches a wheel came off while the robot was operating.\n\nBecause the match was still ongoing we had to react quickly.\n\nWe immediately divided responsibilities. One teammate stabilized the robot while another teammate and I replaced the wheel assembly.\n\nEven though it was stressful the team stayed calm and focused on solving the problem step by step.\n\nThat experience taught me how important communication and coordination are when unexpected problems happen in a team environment.",
            ],
        },
        {
            "id": 3,
            "title": "3. Product oriented engineering mindset",
            "variants": [
                "10s\n\nI enjoy working on problems where system architecture and user experience intersect.",
                "20s\n\nI'm particularly interested in building systems where architecture and product usability come together.\n\nI like thinking about both how the system works internally and how people interact with it.",
                "40s\n\nI tend to approach engineering problems from both a system design and product perspective.\n\nWhen designing systems I think not only about the internal architecture but also how users will understand and interact with it.\n\nThis perspective influenced projects like Spiral, where I tried to turn an abstract system idea into something people could actually use.",
                "90s\n\nOne thing that shapes how I approach engineering is thinking about both system architecture and product usability.\n\nI enjoy working on problems where the internal structure of the system and the user experience influence each other.\n\nWhen building the Spiral prototype I wasn't only focused on the execution model itself.\n\nI also designed an interface that would make the system understandable for users who didn't know the internal implementation.\n\nThat combination of system design and product thinking is something I find very interesting and it's the kind of work I hope to continue doing as an engineer.",
            ],
        },
    ]

    doc = Document()
    doc.add_heading("CS — Core Story Pack", 0)
    doc.add_paragraph()
    for q in cs_core:
        add_question_block(doc, q, "cs-core")
    doc.save("cs-core.docx")
    print("Created cs-core.docx")

    # --- en-interview ---
    en_interview = [
        {
            "id": 1,
            "title": "1. tell me about yourself",
            "variants": [
                "Hi, my name is Juntao Xue. I'm currently based in Tokyo and finishing a degree in History at Rikkyo University, where I focus on institutional systems and how complex structures evolve over time.\n\nAlongside my studies, I've been working on several engineering projects, mainly around system design and product prototyping.\n\nMy main project is called Spiral, which explores how event histories can be represented as structured systems instead of simple mutable state.\n\nRecently I've been focusing on building full-stack prototypes with React, Node, and PostgreSQL, and turning research-style ideas into usable software.",
                "I'm Juntao Xue, a product-oriented engineer currently based in Tokyo.\n\nMy background is actually in history, but my work has gradually moved toward building systems and developer tools. I'm especially interested in how complex systems manage events and state.\n\nMost of my recent work has been around a project called Spiral, where I designed an execution model based on append-only event histories.\n\nIn practice I spend most of my time building full-stack prototypes and experimenting with system architecture and product interaction design.",
            ],
        },
        {
            "id": 2,
            "title": "2. why do you want to work here",
            "variants": [
                "What interests me about your company is the focus on building real products rather than just isolated technical components.\n\nI enjoy environments where engineers are involved not only in implementation but also in shaping how the system and the product evolve.\n\nFrom what I've seen, your team seems to value thoughtful system design and practical engineering, which aligns well with how I like to work.",
                "One reason I'm interested in this role is that your work sits at the intersection of product design and system engineering.\n\nIn my own projects I've spent a lot of time thinking about how architecture decisions affect usability and developer workflows.\n\nSo I'm particularly excited about opportunities where engineering decisions directly influence the user-facing product.",
            ],
        },
        {
            "id": 3,
            "title": "3. describe a challenging technical problem",
            "variants": [
                "One challenging problem I worked on was designing the execution model for the Spiral system.\n\nOriginally the system relied on traditional mutable state, but this made it difficult to reason about how the system reached a particular state.\n\nTo address that, I redesigned the core to use an append-only event graph. Instead of overwriting state, every change becomes a new event, and the system computes a recent view from the causal relationships between events.\n\nThe main challenge was designing rules that kept the model consistent while still making it usable in practice.",
                "A technical challenge I faced was figuring out how to represent complex event histories in a way that was both traceable and efficient.\n\nEarly versions of my project used a more conventional state-update approach, but debugging and reasoning about the system became difficult.\n\nI eventually switched to an append-only event model, where state is derived rather than mutated.\n\nThat change required redesigning several parts of the system, but it made the behavior much easier to understand and debug.",
            ],
        },
        {
            "id": 4,
            "title": "4. tell me about a project you're proud of",
            "variants": [
                "One project I'm particularly proud of is the Spiral prototype platform.\n\nThe goal was to take a fairly abstract system design idea and turn it into something people could actually interact with.\n\nI built a full-stack platform where users can define fragments and observe how event relationships evolve over time.\n\nThe most interesting part for me was designing the interface so that users could understand what the system was doing without needing to read the underlying code.",
                "A project I'm proud of is an experimental system I built called Spiral.\n\nThe idea was to explore a different way of handling event histories in software systems.\n\nI implemented the core execution model and then built a prototype platform around it so users could interact with the system through a structured editor.\n\nIt was rewarding because it combined system architecture, implementation, and product design in a single project.",
            ],
        },
        {
            "id": 5,
            "title": "5. what are your strengths as an engineer",
            "variants": [
                "One of my strengths is being able to think about systems at a structural level.\n\nWhen I work on a problem, I usually try to understand the underlying model first before jumping into implementation.\n\nThat approach has helped me when designing systems like Spiral, where the execution model and data structure have a big impact on how the system behaves.\n\nI also enjoy bridging the gap between architecture decisions and the user-facing product.",
                "I think one of my strengths is turning abstract ideas into working prototypes.\n\nI often start with a conceptual problem, design a simple execution model around it, and then build a full-stack prototype to test whether the idea actually works in practice.\n\nThat process helps me move quickly from exploration to something tangible that people can interact with.",
            ],
        },
        {
            "id": 6,
            "title": "6. describe a technical project or experience",
            "variants": [
                "One project I spent a lot of time on recently is a system called Spiral.\n\nThe original idea was to explore how complex event histories could be represented in software systems in a more transparent way.\n\nInstead of mutating state directly, the system stores events in an append-only structure and derives the current state from their relationships.\n\nI implemented the core execution model and later built a prototype platform so users could interact with it through a structured editor.\n\nWhat I enjoyed most about the project was connecting the system design with a usable product interface.",
                "A project that influenced how I think about system design is something I built called Spiral.\n\nThe main goal was to experiment with representing system behavior through explicit event histories rather than traditional mutable state.\n\nI designed a small execution model for handling those events and implemented a prototype to see how the idea worked in practice.\n\nThrough that process I learned a lot about how architectural choices affect usability and debugging.",
                "One technical project I worked on explored how software systems manage sequences of events.\n\nI built an experimental system where events are stored in an append-only graph, and the system computes a meaningful \"recent state\" from their causal relationships.\n\nThis approach made it easier to understand how the system reached a particular state.\n\nThe project eventually became the basis for a small prototype platform where users could interact with the model through a visual editor.",
                "Recently I've been working on a project that focuses on system architecture and execution models.\n\nThe idea was to experiment with a different way of representing system state by using event histories instead of destructive updates.\n\nI implemented the core logic and then built a small platform around it to test how users could interact with the system.\n\nIt was interesting because it combined backend system design with user-facing product considerations.",
                "One thing I enjoy doing is taking abstract ideas and turning them into working prototypes.\n\nFor example, in one project I explored how event-driven systems could be represented more explicitly.\n\nI designed a small execution model, implemented it, and then built a simple interface around it so people could experiment with the concept.\n\nThat process helped me understand how system architecture decisions affect the usability of the final product.",
            ],
        },
    ]

    doc = Document()
    doc.add_heading("EI — En Interview Pack", 0)
    doc.add_paragraph()
    for q in en_interview:
        add_question_block(doc, q, "en-interview")
    doc.save("en-interview.docx")
    print("Created en-interview.docx")

if __name__ == "__main__":
    main()
